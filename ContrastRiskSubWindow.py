                                                ################################################
                                                #      setting things for CIN RISK Window      #
                                                ################################################
import os
from docx import Document
from nephrology_equations_module import contrast_risk
from PySide6.QtWidgets import QWidget, QLineEdit, QGridLayout, QLabel, QComboBox, QWidget, QPushButton, QCheckBox


class ContrastRiskSubWindow(QWidget):
    def __init__(self):
        super().__init__()
        def check_boxes_logic(status):
            if status.isChecked():
                return "y"
            else:
                return "n"

        def calculate_CIN_risk():
            score,risk,dialysis=contrast_risk(
                egfr=float(egfrLine.text()),
                age=float(ageLine.text()),
                contrast=float(contrastLine.text()),
                hf=check_boxes_logic(hfcbx),
                baloon=check_boxes_logic(balooncbx),
                hypo_tension=check_boxes_logic(shockcbx),
                dm=check_boxes_logic(diabeticcbx),
                anemia=check_boxes_logic(anemialcbx)
            )
            risklbl.setText("Risk for CIN "+str(risk)+"%")
            dialysislbl.setText("Risk for Dialysis "+str(dialysis)+"%")
        def handel_docx(replacements):
            print(replacements)
            doc = Document("pciForm.docx")
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    for key,value in replacements.items():
                        if key in run.text:

                            run.text = run.text.replace(key, value)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                print(run.text)
                                for key,value in replacements.items():

                                    if key in run.text:
                                        print(run.text)
                                        run.text = run.text.replace(key, value)
            doc.save("temppci.docx")
            os.startfile("temppci.docx")
        def generate_docx_report():
            score, risk, dialysis = contrast_risk(
                egfr=float(egfrLine.text()),
                age=float(ageLine.text()),
                contrast=float(contrastLine.text()),
                hf=check_boxes_logic(hfcbx),
                baloon=check_boxes_logic(balooncbx),
                hypo_tension=check_boxes_logic(shockcbx),
                dm=check_boxes_logic(diabeticcbx),
                anemia=check_boxes_logic(anemialcbx)
            )
            replacements = {}
            replacements['putgfr'] = egfrLine.text()
            replacements['putcinhere'] = str(risk)
            replacements['putdiaysishere'] = str(dialysis)
            replacements['putcontrasthere'] = contrastLine.text()
            handel_docx(replacements)
            risklbl.setText("Risk for CIN " + str(risk) + "%")
            dialysislbl.setText("Risk for Dialysis " + str(dialysis) + "%")
        CINLayout = QGridLayout(self)
        self.setLayout(CINLayout)
        self.setWindowTitle("Contrast Risk Calculator")

        egfrlbl=QLabel("eGFR")
        CINLayout.addWidget(egfrlbl,1,1)
        egfrLine=QLineEdit()
        CINLayout.addWidget(egfrLine,1,2)

        agelbl=QLabel("Age")
        CINLayout.addWidget(agelbl,2,1)
        ageLine=QLineEdit()
        CINLayout.addWidget(ageLine,2,2)

        contrastlbl=QLabel("Contrast volume")
        CINLayout.addWidget(contrastlbl,3,1)
        contrastLine=QLineEdit()
        CINLayout.addWidget(contrastLine,3,2)

        hflbl=QLabel("Heart Failure")
        CINLayout.addWidget(hflbl,4,1)
        hfcbx=QCheckBox("Check if present")
        CINLayout.addWidget(hfcbx,4,2)

        baloonlbl=QLabel("Need Baloon")
        CINLayout.addWidget(baloonlbl,5,1)
        balooncbx=QCheckBox("Check if present")
        CINLayout.addWidget(balooncbx,5,2)

        shocklbl=QLabel("Hypotension")
        CINLayout.addWidget(shocklbl,6,1)
        shockcbx=QCheckBox("Check if present")
        CINLayout.addWidget(shockcbx,6,2)

        diabeticlbl=QLabel("Diabetic")
        CINLayout.addWidget(diabeticlbl,7,1)
        diabeticcbx=QCheckBox("Check if present")
        CINLayout.addWidget(diabeticcbx,7,2)

        anemialbl=QLabel("Anemia")
        CINLayout.addWidget(anemialbl,8,1)
        anemialcbx=QCheckBox("Check if present")
        CINLayout.addWidget(anemialcbx,8,2)

        risklbl=QLabel("Risk")
        CINLayout.addWidget(risklbl,9,1)
        dialysislbl=QLabel("Dialysis")
        CINLayout.addWidget(dialysislbl,9,2)

        calculatebtn=QPushButton("Calculate")
        CINLayout.addWidget(calculatebtn,10,1,1,2)
        calculatebtn.clicked.connect(calculate_CIN_risk)

        generate_report_btn=QPushButton("Generate Report")
        CINLayout.addWidget(generate_report_btn,11,1,1,2)
        generate_report_btn.clicked.connect(generate_docx_report)

        print(self.width())