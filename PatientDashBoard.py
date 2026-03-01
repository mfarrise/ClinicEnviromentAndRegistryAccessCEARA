import os
import re
import sys
import json
from datetime import datetime
from PySide6.QtWidgets import QWidget, QGridLayout, QPushButton, QFileDialog, QLineEdit, QApplication, QTextEdit, \
    QLabel, QComboBox, QMenuBar
from docx import Document
from SharedWidgetsPyside6 import show_warning
from PySide6.QtGui import QTextCharFormat, QColor, QFont, QIntValidator, Qt


class PatientDashBoard(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("title")

        self.settings={}
        quadrant_width=540
        quadrant_height=300
        with open("PatientDashBoard_settings.json","r") as file:
            self.settings=json.load(file)

        iraq_governorates = [
            "Baghdad",
            "Basra",
            "Nineveh",
            "Erbil",
            "Sulaymaniyah",
            "Duhok",
            "Kirkuk",
            "Anbar",
            "Babil",
            "Karbala",
            "Najaf",
            "Wasit",
            "Diyala",
            "Salah al-Din",
            "Maysan",
            "Dhi Qar",
            "Al-Muthanna",
            "Al-Qadisiyyah"
        ]
        ##############################################################################
        #setting main lay out that will contain 4 Qgridlayouts each for each quadrant#
        ##############################################################################
        main_layout = QGridLayout(self)
        self.setLayout(main_layout)
        self.move(0, 0)

        menu_bar=QMenuBar()
        file_menu=menu_bar.addMenu("File")
        set_directory_action=file_menu.addAction("Set Directory")
        set_directory_action.triggered.connect(self.set_directory_for_patient_registry_in_setting_json)
        main_layout.addWidget(menu_bar)


        ######################################################################
        #setting up the ULQ which is the patient old history and demographics#
        ######################################################################
        # region
        self.patient_demo_and_old_data_widget=QWidget()
        self.patient_demo_and_old_data_layout = QGridLayout(self.patient_demo_and_old_data_widget)
        self.patient_demo_and_old_data_widget.setLayout(self.patient_demo_and_old_data_layout)
        self.set_quadrants_size(self.patient_demo_and_old_data_widget,quadrant_width,quadrant_height)

        self.patient_id_line_edit=QLineEdit()
        self.patient_id_line_edit.setPlaceholderText("Patient ID")
        self.patient_id_line_edit.setReadOnly(True)
        self.patient_id_line_edit.setFocusPolicy(Qt.ClickFocus)

        self.patient_name_edit = QLineEdit()
        self.patient_name_edit.setPlaceholderText("Patient Name")

        self.patient_DOB_edit = QLineEdit()
        self.patient_DOB_edit.setPlaceholderText("DOB")
        now=datetime.now().year
        self.patient_DOB_edit.setValidator(QIntValidator(1900,now))

        self.patient_gender_combo=QComboBox()
        self.patient_gender_combo.addItems(["Male","Female"])

        self.patient_marital_combo=QComboBox()
        self.patient_marital_combo.addItems(["Single","Married","Divorced","Widowed"])

        self.patient_education_combo=QComboBox()
        self.patient_education_combo.addItems(["University","PostGraduate","Institute","Secondary","Primary","None"])

        self.patient_job_edit = QLineEdit()
        self.patient_job_edit.setPlaceholderText("Job")

        self.patient_governorates_combo = QComboBox()
        self.patient_governorates_combo.addItems(iraq_governorates)

        self.patient_residence_free_form_edit=QLineEdit()
        self.patient_residence_free_form_edit.setPlaceholderText("residence")

        self.patient_residence_type_combo = QComboBox()
        self.patient_residence_type_combo.addItems(["center","periphery","rural"])



        self.patient_demo_and_old_data_layout.addWidget(self.patient_id_line_edit,0,0)
        self.patient_demo_and_old_data_layout.addWidget(self.patient_name_edit,0,1)
        self.patient_demo_and_old_data_layout.addWidget(self.patient_DOB_edit,0,2)
        self.patient_demo_and_old_data_layout.addWidget(self.patient_gender_combo,0,3)
        self.patient_demo_and_old_data_layout.addWidget(self.patient_marital_combo,0,4)
        self.patient_demo_and_old_data_layout.addWidget(self.patient_education_combo,1,0)
        self.patient_demo_and_old_data_layout.addWidget(self.patient_job_edit,1,1)
        self.patient_demo_and_old_data_layout.addWidget(self.patient_governorates_combo, 1, 2)
        self.patient_demo_and_old_data_layout.addWidget(self.patient_residence_free_form_edit, 1, 3)
        self.patient_demo_and_old_data_layout.addWidget(self.patient_residence_type_combo,1,4)









        #previous history
        # label
        self.previous_history_label = QLabel()
        self.previous_history_label.setText("Previous History")
        self.patient_demo_and_old_data_layout.addWidget(self.previous_history_label,2,0)

        # previous history
        # text edit
        self.previous_history_edit = QTextEdit()
        self.previous_history_edit.setReadOnly(True)
        self.patient_demo_and_old_data_layout.addWidget(self.previous_history_edit, 3, 0,1,5)

        main_layout.addWidget(self.patient_demo_and_old_data_widget,1,0)
        # endregion
        ##########################################################################
        # setting up the LLQ which is the patient current history and examination#
        ##########################################################################
        #region
        # creat 2nd degree nested two layouts and widget inside the left lower quadrant llQ layout
        self.llql_widget=QWidget()
        self.set_quadrants_size(self.llql_widget, int(quadrant_width / 2), quadrant_height)
        self.llql_layout = QGridLayout()
        self.llql_widget.setLayout(self.llql_layout)

        self.llqr_widget=QWidget()
        self.set_quadrants_size(self.llqr_widget, int(quadrant_width / 2), quadrant_height)
        self.llqr_layout = QGridLayout()
        self.llqr_widget.setLayout(self.llqr_layout)


        self.patient_today_clinical_widget=QWidget()
        self.patient_today_clinical_layout = QGridLayout()
        self.patient_today_clinical_widget.setLayout(self.patient_today_clinical_layout)
        self.set_quadrants_size(self.patient_today_clinical_widget,quadrant_width,quadrant_height)

        self.patient_today_clinical_layout.addWidget(self.llql_widget,0,0)
        self.patient_today_clinical_layout.addWidget(self.llqr_widget, 0, 1)

        #Today history
        #label
        today_history_label = QLabel(self)
        today_history_label.setText("Today clinical")
        self.llql_layout.addWidget(today_history_label, 0, 0)

        # Today history
        #text Edit
        today_history_edit = QTextEdit(self)
        today_history_edit.setReadOnly(False)
        self.llql_layout.addWidget(today_history_edit, 1, 0)
        self.symptoms_line_edit=[]

        for t in range(1,4):
            for i in range(0,7):
                line_edit=QLineEdit()
                line_edit.setPlaceholderText("entry")
                self.llqr_layout.addWidget(line_edit, i, t)
                self.symptoms_line_edit.append(line_edit)
        # print(len(self.symptoms_line_edit))


        main_layout.addWidget(self.patient_today_clinical_widget,2,0)
        #endregion
        ##########################################################################
        # setting up the RUQ which is the patient current history and examination#
        ##########################################################################
        #region
        patient_today_ix_medication_widget=QWidget(self)
        patient_today_ix_medication_layout = QGridLayout(patient_today_ix_medication_widget)
        patient_today_ix_medication_widget.setLayout(patient_today_ix_medication_layout)
        self.set_quadrants_size(patient_today_ix_medication_widget,quadrant_width,quadrant_height)

        # today investigations
        # label
        today_investigations_label = QLabel(self)
        today_investigations_label.setText("Today Investigations")
        patient_today_ix_medication_layout.addWidget(today_investigations_label, 1, 1)

        # today investigations
        # text edit
        self.today_investigations_edit = QTextEdit(self)
        self.today_investigations_edit.setReadOnly(False)
        patient_today_ix_medication_layout.addWidget(self.today_investigations_edit, 2, 1, 1, 1)

        # Last medications == today modeications
        # label
        today_medications_label = QLabel(self)
        today_medications_label.setText("Today Medications")
        patient_today_ix_medication_layout.addWidget(today_medications_label, 1, 2)

        # today medications
        # text edit
        self.today_medications_edit = QTextEdit(self)
        self.today_medications_edit.setReadOnly(False)
        patient_today_ix_medication_layout.addWidget(self.today_medications_edit, 2, 2, 1, 1)

        # today investigation intelligent
        # line edit
        self.today_investigations_intellisense_line = QLineEdit(self)
        self.today_investigations_intellisense_line.setPlaceholderText("Intelligent Fill")
        self.today_investigations_intellisense_line.setReadOnly(False)
        patient_today_ix_medication_layout.addWidget(self.today_investigations_intellisense_line, 3, 1, 1, 1)
        self.today_investigations_intellisense_line.returnPressed.connect(self.investigation_intellisense)
        # today Medication intelligent
        # line edit
        self.today_medication_intellisense_line = QLineEdit(self)
        self.today_medication_intellisense_line.setPlaceholderText("Intelligent Fill")
        self.today_medication_intellisense_line.setReadOnly(False)
        patient_today_ix_medication_layout.addWidget(self.today_medication_intellisense_line, 3, 2, 1, 1)
        self.today_medication_intellisense_line.returnPressed.connect(self.drug_intellisense)
        main_layout.addWidget(patient_today_ix_medication_widget,1,1)

        #endregion
        ##########################################################################
        # setting up the RUQ which is the patient current history and examination#
        ##########################################################################
        ##########################################
        #starting writing th engine of the widget#
        ##########################################
        # region

        def load_patient_data():
            patient_file_path = self.open_path_dialog()
            sep = r"[\/\.\-ظ]"

            date_pattern = rf"\b\d{{1,2}}{sep}\d{{1,2}}{sep}\d{{2,4}}\b|\b\d{{4}}{sep}\d{{1,2}}{sep}\d{{1,2}}\b"

            self.previous_history_edit.clear()
            cursor = self.previous_history_edit.textCursor()
            if patient_file_path:
                doc = Document(patient_file_path)
                # in the following code
                # splitext() returns a tuple:
                # ("README", ".md")
                # [0]takes the first part(without extension)
                self.patient_name_edit.setText(os.path.splitext(os.path.basename(patient_file_path))[0])
                # old_history=""
                normal_format = QTextCharFormat()

                date_format = QTextCharFormat()
                date_format.setForeground(QColor("orange"))
                date_format.setFontWeight(QFont.Bold)
                cr_format = QTextCharFormat()
                cr_format.setForeground(QColor("cyan"))
                cr_format.setFontWeight(QFont.Bold)
                aliases_list_for_oldhistory=[]

                for paragraph in doc.paragraphs:
                    text = paragraph.text

                    parts = re.split(f"({date_pattern})", text)

                    for part in parts:
                        if re.fullmatch(date_pattern, part):
                            cursor.setCharFormat(date_format)
                            cursor.insertText(part)
                            cursor.setCharFormat(normal_format)  # 🔥 RESET i know its redundant due to reset in else but for future unforseen changes
                        elif re.search(r"\bcr|creat|creatinine|urea|pus|gue\b", part,re.IGNORECASE):
                            cursor.setCharFormat(cr_format)
                            cursor.insertText(part)
                            cursor.setCharFormat(normal_format)
                        else:
                            cursor.setCharFormat(normal_format)
                            cursor.insertText(part)

                    cursor.insertText("\n")
                # previous_history_edit.setText(old_history)

        def update_patient_data():

            doc = Document()
            demographic_table=doc.add_table(2,5)

            demographic_table.cell(0,0).text="Name"
            if self.patient_name_edit.text() != "":
                demographic_table.cell(1,0).text=self.patient_name_edit.text()

            demographic_table.cell(0,1).text="DOB"
            if self.patient_name_edit.text() != "":
                demographic_table.cell(1,1).text=self.patient_DOB_edit.text()

            demographic_table.cell(0,2).text="Age"
            if self.patient_name_edit.text() != "":
                demographic_table.cell(1,1).text=str(datetime.now().year-int(self.patient_DOB_edit.text()))

            demographic_table.cell(0,3).text="Gender"
            demographic_table.cell(1,3).text=self.patient_gender_combo.currentText()

            demographic_table.cell(0,4).text="Residence"
            if self.patient_name_edit.text() != "":
                demographic_table.cell(1,4).text=self.patient_residence_free_form_edit.text()

            doc.save("PatientDashBoard.docx")


        io_widget=QWidget(self)
        io_layout = QGridLayout(io_widget)
        io_widget.setLayout(io_layout)
        self.set_quadrants_size(io_widget,quadrant_width,quadrant_height)

        #load patient button
        load_patient_button=QPushButton("Load Patient")
        io_layout.addWidget(load_patient_button,0,0)
        load_patient_button.clicked.connect(load_patient_data)

        #update new data button
        update_patient_button=QPushButton("Update Patient")
        io_layout.addWidget(update_patient_button,0,1)
        update_patient_button.clicked.connect(update_patient_data)



        main_layout.addWidget(io_widget,2,1)

        #endregion


    def open_path_dialog(self):

        file_path, _ = QFileDialog.getOpenFileName(
            parent=None,
            caption="Select file",
            dir=self.settings["patient_dir"],#settings is a dictionary read from json this dictionary hold the settings including patient files directory
            filter="All Files (*.*);;Word Files (*.docx);;PDF Files (*.pdf)"
        )
        #print(type(file_path),file_path)
        return file_path

    def set_quadrants_size(self,widget, width=450, height=450):
        widget.setMaximumWidth(width)
        widget.setMaximumHeight(height)
        widget.setMinimumHeight(height)
        widget.setMinimumWidth(width)

    def set_directory_for_patient_registry_in_setting_json(self):
        temp_dir_path=self.pick_directory_name()
        if temp_dir_path:#to check if a path was selected
            self.settings["patient_dir"]=temp_dir_path
            with open("PatientDashBoard_settings.json","w") as file:
                json.dump(self.settings,file,indent=4)

    def pick_directory_name(self):
        dir_path = QFileDialog.getExistingDirectory(
            parent=None,
            caption="Select file",
            dir="",

        )

        print(type(dir_path),dir_path)
        return dir_path

    def investigation_intellisense(self):
        unparsed_text=self.today_investigations_intellisense_line.text()
        self.today_investigations_intellisense_line.clear()# clear feild after value is taken in line above
        parsed_text=unparsed_text.split()

        if len(parsed_text) < 2:
            show_warning("Not a valid test result")
            return  # Not enough data resulting list is 1 word
        if len (parsed_text) ==3:# if things like b urea 15 it ill parse it into b.urea 15 so it wont stay in to 3 element list ["b","urea","15"]
            parsed_text=[parsed_text[0]+"."+parsed_text[1],parsed_text[2]]
            print (parsed_text)
        if not unparsed_text:# dont know why but this line is not functioning
            return
        try:
            float(parsed_text[-1])
        except ValueError:
            show_warning("Please enter a numerical value after test name")
            return
        with open("InvestigationDatabase.json","r") as file:
            InvestigationDatabase=json.load(file)

        for case in InvestigationDatabase.values():

            if parsed_text[0] in case["aliases"]:
                self.append_investigation(
                    case["display"],
                    float(parsed_text[-1]),
                    case["high"],
                    case["low"],
                    case["unit"]
                )
    def append_investigation(self,test_name,test_value,upper_limit,lower_limit,unit):
        if test_value > upper_limit:
            self.today_investigations_edit.append(
                f"<span style='color:rgb(200,110,110)' ><b>{test_name+' ' + str(test_value) + ' ' +unit+' (High)'}</b></span>"
            )
        elif test_value < lower_limit:
            self.today_investigations_edit.append(
                f"<span style='color:rgb(200,110,110)' ><b>{test_name + ' ' + str(test_value) + ' ' + unit + ' (Low)'}</b></span>"
            )
        else :
            self.today_investigations_edit.append(
                f"<span ><b>{test_name + ' ' + str(test_value) + ' ' + unit }</b></span>"
            )

    def drug_intellisense(self):#not fully implemented yet there is no intellgence in it now
        self.today_medications_edit.append(self.today_medication_intellisense_line.text())
        self.today_medication_intellisense_line.clear()


if __name__ == '__main__':
    app = QApplication(sys.argv)
    window = PatientDashBoard()
    window.show()
    app.exec()
